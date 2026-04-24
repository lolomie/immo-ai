"""
Word document (.docx) creator for validated exposés.
Design: Anthrazit / Creme / Gold — matching the web frontend.
"""

import io
import logging
import re
from datetime import datetime, timezone
from typing import Optional

from .models import PropertyInput

logger = logging.getLogger(__name__)

_DISCLAIMER = (
    "Dieses Exposé wurde mit Unterstützung von KI-Technologie erstellt und durch "
    "einen menschlichen Makler geprüft. Alle Angaben sind nach bestem Wissen und "
    "Gewissen, jedoch ohne Gewähr. Irrtümer und Zwischenverkauf vorbehalten."
)

# ── Design tokens ──────────────────────────────────────────────────────────────
_GOLD        = (0xB8, 0x97, 0x3A)
_INK         = (0x2B, 0x2B, 0x2B)
_INK2        = (0x4A, 0x4A, 0x4A)
_MUTED       = (0x9E, 0x95, 0x89)
_WHITE       = (0xFF, 0xFF, 0xFF)
_CREAM_HEX   = "F5F0E8"
_BORDER_HEX  = "E5DFD4"
_PAPER_HEX   = "FAFAF7"
_GOLD_HEX    = "B8973A"
_INK_HEX     = "2B2B2B"

_FONT_SERIF  = "Times New Roman"
_FONT_SANS   = "Calibri"


# ── XML helpers ────────────────────────────────────────────────────────────────

def _xml_el(tag):
    from docx.oxml import OxmlElement
    return OxmlElement(tag)

def _qn(tag):
    from docx.oxml.ns import qn
    return qn(tag)

def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = _xml_el("w:shd")
    shd.set(_qn("w:val"), "clear")
    shd.set(_qn("w:color"), "auto")
    shd.set(_qn("w:fill"), hex_color)
    tcPr.append(shd)

def _set_cell_borders(cell, color_hex: str = _BORDER_HEX) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBdr = _xml_el("w:tcBdr")
    for side in ("top", "left", "bottom", "right"):
        el = _xml_el(f"w:{side}")
        el.set(_qn("w:val"), "single")
        el.set(_qn("w:sz"), "4")
        el.set(_qn("w:space"), "0")
        el.set(_qn("w:color"), color_hex)
        tcBdr.append(el)
    tcPr.append(tcBdr)

def _remove_table_borders(table) -> None:
    """Remove all outer/inner borders from a table."""
    from docx.oxml import OxmlElement
    tbl = table._tbl
    tblPr = tbl.find(_qn("w:tblPr"))
    if tblPr is None:
        tblPr = _xml_el("w:tblPr")
        tbl.insert(0, tblPr)
    tblBdr = _xml_el("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = _xml_el(f"w:{side}")
        el.set(_qn("w:val"), "none")
        tblBdr.append(el)
    tblPr.append(tblBdr)

def _set_para_bottom_border(para, color_hex: str, sz: str = "6") -> None:
    pPr = para._p.get_or_add_pPr()
    pBdr = _xml_el("w:pBdr")
    bot = _xml_el("w:bottom")
    bot.set(_qn("w:val"), "single")
    bot.set(_qn("w:sz"), sz)
    bot.set(_qn("w:space"), "1")
    bot.set(_qn("w:color"), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)

def _set_cell_no_padding(cell) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = _xml_el("w:tcMar")
    for side in ("top", "left", "bottom", "right"):
        m = _xml_el(f"w:{side}")
        m.set(_qn("w:w"), "80")
        m.set(_qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)


# ── Markdown parser ────────────────────────────────────────────────────────────

def _add_md_runs(para, text: str, base_size_pt: float, color_rgb: tuple,
                 font: str = _FONT_SANS, base_bold: bool = False) -> None:
    """Split text on **bold** and *italic* markers, add styled runs."""
    from docx.shared import Pt, RGBColor
    parts = re.split(r"(\*\*.*?\*\*|\*[^*].*?[^*]\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = para.add_run(part[1:-1])
            run.italic = True
        else:
            run = para.add_run(part)
            if base_bold:
                run.bold = True
        run.font.name = font
        run.font.size = Pt(base_size_pt)
        run.font.color.rgb = RGBColor(*color_rgb)


def _parse_expose_text(doc, expose_text: str) -> None:
    """
    Render the expose body into the document.
    Handles ## headings, **bold**, *italic*, and plain paragraphs.
    Strips all raw asterisks that aren't part of markup.
    """
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    paragraphs = expose_text.strip().split("\n\n")

    for block in paragraphs:
        block = block.strip()
        if not block:
            continue

        lines = block.splitlines()
        first_line = lines[0].strip()

        # Detect heading (## or ### prefix, or line ending with ':' after **...**)
        heading_match = re.match(r"^(#{1,3})\s+(.*)", first_line)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip().strip("*")
            word_level = min(level + 1, 3)
            h = doc.add_heading("", level=word_level)
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after = Pt(4)
            run = h.add_run(heading_text)
            run.font.name = _FONT_SERIF
            run.font.size = Pt(12 if level > 1 else 13)
            run.font.color.rgb = RGBColor(*_INK)
            run.font.bold = True
            # Add remaining lines as normal paragraph
            rest = "\n".join(lines[1:]).strip()
            if rest:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(6)
                _add_md_runs(p, rest, 10.5, _INK2)
            continue

        # Detect bold-only heading line (standalone **text**)
        bold_heading_match = re.match(r"^\*\*([^*]+)\*\*\s*$", first_line)
        if bold_heading_match:
            h = doc.add_heading("", level=3)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(3)
            run = h.add_run(bold_heading_match.group(1))
            run.font.name = _FONT_SERIF
            run.font.size = Pt(11.5)
            run.font.color.rgb = RGBColor(*_INK)
            run.bold = True
            rest = "\n".join(lines[1:]).strip()
            if rest:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(6)
                _add_md_runs(p, rest, 10.5, _INK2)
            continue

        # Normal paragraph — join lines, parse inline markup
        full_text = " ".join(lines)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(7)
        _add_md_runs(p, full_text, 10.5, _INK2)


# ── Main creator ───────────────────────────────────────────────────────────────

def create_expose_docx_bytes(
    property_data: PropertyInput,
    expose_text: str,
    agent_email: Optional[str] = None,
    target_group: Optional[str] = None,
) -> bytes:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, RGBColor, Inches, Cm
    except ImportError as e:
        raise ImportError("python-docx not installed. Run: pip install python-docx") from e

    doc = Document()

    # ── Page setup ────────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(2.8)

    # Remove default paragraph spacing from Normal style
    doc.styles["Normal"].paragraph_format.space_after  = Pt(0)
    doc.styles["Normal"].paragraph_format.space_before = Pt(0)
    doc.styles["Normal"].font.name  = _FONT_SANS
    doc.styles["Normal"].font.size  = Pt(10.5)
    doc.styles["Normal"].font.color.rgb = RGBColor(*_INK2)

    # ── ① Gold header bar ─────────────────────────────────────────────────────
    header_tbl = doc.add_table(rows=1, cols=2)
    _remove_table_borders(header_tbl)
    header_tbl.autofit = False
    header_tbl.columns[0].width = Cm(14)
    header_tbl.columns[1].width = Cm(3)

    left_cell  = header_tbl.cell(0, 0)
    right_cell = header_tbl.cell(0, 1)
    _set_cell_bg(left_cell, "2B2B2B")
    _set_cell_bg(right_cell, _GOLD_HEX)

    lp = left_cell.paragraphs[0]
    lp.paragraph_format.space_before = Pt(0)
    lp.paragraph_format.space_after  = Pt(0)
    lr = lp.add_run("IMMO AI  —  IMMOBILIEN-EXPOSÉ")
    lr.font.name  = _FONT_SANS
    lr.font.size  = Pt(8)
    lr.font.bold  = True
    lr.font.color.rgb = RGBColor(0xD4, 0xB0, 0x5A)
    lr.font.all_caps  = True

    rp = right_cell.paragraphs[0]
    rp.paragraph_format.space_before = Pt(0)
    rp.paragraph_format.space_after  = Pt(0)
    rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = rp.add_run("✦")
    rr.font.name  = _FONT_SANS
    rr.font.size  = Pt(10)
    rr.font.color.rgb = RGBColor(*_WHITE)

    doc.add_paragraph()

    # ── ② Title block ─────────────────────────────────────────────────────────
    rooms_label = ""
    if property_data.rooms:
        r = str(property_data.rooms)
        rooms_label = r.rstrip("0").rstrip(".") + "-Zimmer-"

    title_text = f"{rooms_label}{property_data.property_type}"
    title_para = doc.add_paragraph()
    title_para.paragraph_format.space_before = Pt(16)
    title_para.paragraph_format.space_after  = Pt(4)
    tr = title_para.add_run(title_text)
    tr.font.name  = _FONT_SERIF
    tr.font.size  = Pt(26)
    tr.font.bold  = True
    tr.font.color.rgb = RGBColor(*_INK)

    subtitle_parts = [
        property_data.address,
        f"{property_data.zip_code} {property_data.city}",
    ]
    if property_data.year_built:
        subtitle_parts.append(f"Bj. {property_data.year_built}")

    sub_para = doc.add_paragraph()
    sub_para.paragraph_format.space_before = Pt(0)
    sub_para.paragraph_format.space_after  = Pt(14)
    sr = sub_para.add_run("  ·  ".join(subtitle_parts))
    sr.font.name  = _FONT_SANS
    sr.font.size  = Pt(10)
    sr.font.color.rgb = RGBColor(*_MUTED)

    # Gold divider below title
    _set_para_bottom_border(sub_para, _GOLD_HEX, sz="12")

    doc.add_paragraph()

    # ── ③ Stats bar ───────────────────────────────────────────────────────────
    stats = [
        ("Fläche",    f"{property_data.size_sqm} m²"      if property_data.size_sqm  else "—"),
        ("Zimmer",    str(property_data.rooms).rstrip("0").rstrip(".") if property_data.rooms else "—"),
        ("Baujahr",   str(property_data.year_built)         if property_data.year_built else "—"),
        ("Energie",   property_data.energy_class             if property_data.energy_class else "—"),
    ]

    stats_tbl = doc.add_table(rows=2, cols=4)
    _remove_table_borders(stats_tbl)
    stats_tbl.autofit = True

    for col_i, (label, value) in enumerate(stats):
        val_cell = stats_tbl.cell(0, col_i)
        lbl_cell = stats_tbl.cell(1, col_i)
        _set_cell_bg(val_cell, _CREAM_HEX)
        _set_cell_bg(lbl_cell, _CREAM_HEX)
        _set_cell_borders(val_cell, _BORDER_HEX)
        _set_cell_borders(lbl_cell, _BORDER_HEX)

        vp = val_cell.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        vp.paragraph_format.space_before = Pt(6)
        vp.paragraph_format.space_after  = Pt(0)
        vr = vp.add_run(value)
        vr.font.name  = _FONT_SERIF
        vr.font.size  = Pt(14)
        vr.font.bold  = True
        vr.font.color.rgb = RGBColor(*_GOLD)

        lp = lbl_cell.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp.paragraph_format.space_before = Pt(0)
        lp.paragraph_format.space_after  = Pt(6)
        lr = lp.add_run(label.upper())
        lr.font.name  = _FONT_SANS
        lr.font.size  = Pt(6.5)
        lr.font.bold  = True
        lr.font.color.rgb = RGBColor(*_MUTED)
        lr.font.all_caps  = True

    doc.add_paragraph()

    # ── ④ Details table ───────────────────────────────────────────────────────
    detail_rows = [
        ("Objektart",    property_data.property_type),
        ("Adresse",      f"{property_data.address}, {property_data.zip_code} {property_data.city}"),
        ("Wohnfläche",   f"{property_data.size_sqm} m²"),
        ("Zimmer",       str(property_data.rooms).rstrip("0").rstrip(".")),
    ]
    if property_data.year_built:
        detail_rows.append(("Baujahr",      str(property_data.year_built)))
    if property_data.energy_class:
        detail_rows.append(("Energieklasse", property_data.energy_class))
    if property_data.purchase_price:
        detail_rows.append(("Kaufpreis",    f"{property_data.purchase_price:,.0f} €".replace(",", ".")))
    if property_data.monthly_rent:
        detail_rows.append(("Kaltmiete",    f"{property_data.monthly_rent:,.0f} €/Monat".replace(",", ".")))
    if property_data.features:
        detail_rows.append(("Ausstattung",  "  ·  ".join(property_data.features)))
    if target_group:
        detail_rows.append(("Zielgruppe",   target_group))

    det_tbl = doc.add_table(rows=len(detail_rows), cols=2)
    _remove_table_borders(det_tbl)

    for i, (label, value) in enumerate(detail_rows):
        lbl_cell = det_tbl.cell(i, 0)
        val_cell = det_tbl.cell(i, 1)
        _set_cell_bg(lbl_cell, _CREAM_HEX)
        _set_cell_borders(lbl_cell, _BORDER_HEX)
        _set_cell_borders(val_cell, _BORDER_HEX)

        lp = lbl_cell.paragraphs[0]
        lp.paragraph_format.space_before = Pt(2)
        lp.paragraph_format.space_after  = Pt(2)
        lr = lp.add_run(label)
        lr.font.name  = _FONT_SANS
        lr.font.size  = Pt(9)
        lr.font.bold  = True
        lr.font.color.rgb = RGBColor(*_INK2)

        vp = val_cell.paragraphs[0]
        vp.paragraph_format.space_before = Pt(2)
        vp.paragraph_format.space_after  = Pt(2)
        vr = vp.add_run(value)
        vr.font.name  = _FONT_SANS
        vr.font.size  = Pt(9)
        vr.font.color.rgb = RGBColor(*_INK2)

    doc.add_paragraph()

    # ── ⑤ Section heading: Objektbeschreibung ─────────────────────────────────
    sec_para = doc.add_paragraph()
    sec_para.paragraph_format.space_before = Pt(16)
    sec_para.paragraph_format.space_after  = Pt(10)
    _set_para_bottom_border(sec_para, _GOLD_HEX, sz="6")
    sec_run = sec_para.add_run("OBJEKTBESCHREIBUNG")
    sec_run.font.name  = _FONT_SANS
    sec_run.font.size  = Pt(7.5)
    sec_run.font.bold  = True
    sec_run.font.all_caps  = True
    sec_run.font.color.rgb = RGBColor(*_GOLD)

    # ── ⑥ Exposé body text ────────────────────────────────────────────────────
    _parse_expose_text(doc, expose_text)

    # ── ⑦ Footer ──────────────────────────────────────────────────────────────
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(18)
    sep.paragraph_format.space_after  = Pt(10)
    _set_para_bottom_border(sep, _BORDER_HEX, sz="4")

    meta_lines = [
        f"Erstellt am {datetime.now(timezone.utc).strftime('%d.%m.%Y')}",
        f"Objekt-ID: {property_data.property_id}",
    ]
    if agent_email:
        meta_lines.append(f"Zuständig: {agent_email}")

    meta_para = doc.add_paragraph()
    meta_para.paragraph_format.space_before = Pt(0)
    meta_para.paragraph_format.space_after  = Pt(6)
    mr = meta_para.add_run("  ·  ".join(meta_lines))
    mr.font.name  = _FONT_SANS
    mr.font.size  = Pt(8)
    mr.font.color.rgb = RGBColor(*_MUTED)

    disc_para = doc.add_paragraph()
    disc_para.paragraph_format.space_before = Pt(0)
    disc_para.paragraph_format.space_after  = Pt(0)
    dr = disc_para.add_run(_DISCLAIMER)
    dr.font.name   = _FONT_SANS
    dr.font.size   = Pt(7.5)
    dr.font.italic = True
    dr.font.color.rgb = RGBColor(*_MUTED)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def expose_filename(property_data: PropertyInput) -> str:
    safe_id  = property_data.property_id.replace("/", "-").replace(" ", "_")
    date_str = datetime.now(timezone.utc).strftime("%Y%m%d")
    return f"Expose_{safe_id}_{date_str}.docx"
